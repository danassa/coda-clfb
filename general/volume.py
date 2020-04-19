from docx import Document
import logging
from general.constants import *
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class Volume:
    def __init__(self, chars, start, end, fst_c_index_in_v, lst_c_index_in_v, first_paragraph=None, last_paragraph=None):
        """
        :param chars: total count of characters in all this volume's paragraphs
        :param start: index of the book's paragraphs which will be the FIRST in this volume
        :param end: index of the book's paragraphs which will be the LAST in this volume
        :param fst_c_index_in_v: the index of the FIRST chapter in this volume, in the book's Chapters field
        :param lst_c_index_in_v: the index of the LAST chapter in this volume, according to the book's Chapters field
        :param first_paragraph: indicate if this volume STARTS in the middle of a chapter. None or a string.
        :param last_paragraph: indicate if this volume ENDS in the middle of a chapter. None or a string.
        """
        self.start_index = start
        self.end_index = end
        self.chars = chars
        self.first_paragraph = first_paragraph
        self.last_paragraph = last_paragraph
        self.first_chapter_index = fst_c_index_in_v
        self.last_chapter_index = lst_c_index_in_v
        self.elements = None
        self.doc = None
        self.volume_num = None
        self.font = None
        self.style = None
        logging.debug("Volume initialized with paragraphs: {} to {}, chapters {} to {}, with first paragraph: '{}' and last: '{}'"
                      .format(self.start_index, self.end_index, self.first_chapter_index, self.last_chapter_index, self.first_paragraph, self.last_paragraph))

    def update_first_chapter(self, chars_added, new_first_paragraph):
        self.start_index = new_first_paragraph
        self.chars = self.chars + chars_added
        self.first_paragraph = None

    def update_last_chapter(self, chars_removed, new_last_paragraph):
        self.end_index = new_last_paragraph
        self.chars = self.chars - chars_removed
        self.last_paragraph = None

    def create_doc(self, doc, elements, gate_index, volume_num, last_volume_num):
        self.doc = doc
        self.font = self.doc.paragraphs[gate_index - 1].runs[-1].font
        self.style = self.doc.paragraphs[gate_index - 1].style
        self.elements = elements
        self.volume_num = volume_num
        self.remove_redundant_paragraphs(gate_index)
        self.delete_chapter_beginning_markers()
        self.add_split_chapter_notes(gate_index)
        if volume_num == last_volume_num:
            self.add_book_suffix()
        else:
            self.add_volume_suffix()
        self.update_first_page(last_volume_num)

    def remove_redundant_paragraphs(self, gate_index):
        last_index_in_book = len(self.elements) - 1
        for p in range(last_index_in_book, self.end_index, -1):
            self.elements[p].delete()
        for p in range(self.start_index - 1, gate_index, -1):
            self.elements[p].delete()
        logging.debug("deleting paragraphs {}-{} and {}-{}"
                      .format(self.end_index, last_index_in_book, gate_index, self.start_index))

    def delete_chapter_beginning_markers(self):
        for p in self.elements:
            if p.is_paragraph and p.block.text == CHAPTER_INDICATOR:
                p.block.text = '' #todo delete instead?

    def add_split_chapter_notes(self, gate_index):
        if self.first_paragraph is not None:
            p = self.doc.paragraphs[gate_index + 1].insert_paragraph_before(self.first_paragraph, self.style)
            self.set_style(p)
            self.doc.paragraphs[gate_index + 2].insert_paragraph_before('')
        if self.last_paragraph is not None:
            self.doc.add_paragraph('')
            p = self.doc.add_paragraph(self.last_paragraph, self.style)
            self.set_style(p)

    def add_volume_suffix(self):
        volume_suffix = '{}{}'.format(END, VOLUMES_NAMES_1[self.volume_num - 1])
        self.add_ending("\n\n{}".format(volume_suffix))

    def add_book_suffix(self):
        last_paragraph = self.find_last_non_empty_paragraph_index()
        if self.doc.paragraphs[last_paragraph].text != BOOK_SUFFIX:
            self.add_ending("\n\n{}".format(BOOK_SUFFIX))
            logging.debug("last paragraph in the book was not '{}', adding the phrase".format(BOOK_SUFFIX))

    def add_ending(self, text):
        p = self.doc.add_paragraph(text, self.style)
        self.set_style(p)
        # todo rtl

    def set_style(self, paragraph):
        paragraph.runs[0].font.size = self.font.size
        paragraph.runs[0].font.rtl = self.font.rtl
        paragraph.runs[0].font.name = self.font.name
        paragraph.runs[-1].font.size = self.font.size
        paragraph.runs[-1].font.rtl = self.font.rtl
        paragraph.runs[-1].font.name = self.font.name
        paragraph.style = self.style
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    def update_first_page(self, last_volume_num):
        index = None
        for i in range(30, 1, -1):
            if self.doc.paragraphs[i].text == FIRST_PAGE_VOLUMES_PARAGRAPH:
                index = i
        if index is not None:
            self.update_text(index, VOLUMES_NAMES_1[self.volume_num - 1])
            self.update_text(index - 1, VOLUMES_NAMES_2[last_volume_num - 1])
        else:
            logging.error("couldn't locate the phrase {} in the first 30 paragraphs. "
                          "cannot update the first page with number of volumes".format(FIRST_PAGE_VOLUMES_PARAGRAPH))

    def update_text(self, index, new_text):
        runs = self.doc.paragraphs[index].runs
        for r in runs:
            r.clear()
        runs[0].text = new_text

    def find_last_non_empty_paragraph_index(self):
        paragraph = None
        index = len(self.doc.paragraphs) - 1
        while paragraph is None and index >= 0:
            if self.doc.paragraphs[index].text != '':
                paragraph = index
            index -= 1
        return paragraph
