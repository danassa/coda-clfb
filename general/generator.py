import logging
from general.constants import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from general.element import Element
import platform


class DocxGenerator:

    def __init__(self, doc, elements, gate_index, volume, volume_index, volumes_count):
        self.doc = doc
        self.elements = elements
        self.gate_index = gate_index
        self.volume = volume
        self.volume_index = volume_index
        self.volumes_count = volumes_count
        self.default_font = self.doc.paragraphs[self.gate_index - 1].runs[-1].font

    def format_volume(self):
        self.remove_redundant_paragraphs(self.elements)
        self.delete_chapter_beginning_markers(self.elements, CHAPTER_INDICATOR)
        self.add_split_chapter_notes()
        if self.volume_index == self.volumes_count:
            self.add_book_suffix()
        else:
            self.add_volume_suffix()
        self.update_first_page()

    def remove_redundant_paragraphs(self, elements):
        last_index_in_book = len(elements) - 1
        for p in range(last_index_in_book, self.volume.end_index, -1):
            elements[p].delete()
        for p in range(self.volume.start_index - 1, self.gate_index, -1):
            elements[p].delete()
        logging.debug("deleting paragraphs {}-{} and {}-{}"
                      .format(self.volume.end_index, last_index_in_book, self.gate_index, self.volume.start_index))

    def add_split_chapter_notes(self):
        if self.volume.first_paragraph is not None:
            p = self.doc.paragraphs[self.gate_index + 1].insert_paragraph_before(self.volume.first_paragraph)
            self.set_style(p)
            self.doc.paragraphs[self.gate_index + 2].insert_paragraph_before('')
        if self.volume.last_paragraph is not None:
            self.doc.add_paragraph('')
            p = self.doc.add_paragraph(self.volume.last_paragraph)
            self.set_style(p)

    def add_volume_suffix(self):
        volume_suffix = '{}{}'.format(END, VOLUMES_NAMES_1[self.volume_index - 1])
        self.add_ending("\n\n{}".format(volume_suffix))

    def add_book_suffix(self):
        last_paragraph = None
        index = len(self.doc.paragraphs) - 1
        while last_paragraph is None and index >= 0:
            if self.doc.paragraphs[index].text != '':
                last_paragraph = index
            index -= 1
        if self.doc.paragraphs[last_paragraph].text != BOOK_SUFFIX:
            self.add_ending("\n\n{}".format(BOOK_SUFFIX))
            logging.debug("last paragraph in the book was not '{}', adding the phrase".format(BOOK_SUFFIX))

    def add_ending(self, text):
        p = self.doc.add_paragraph(text)
        self.set_style(p)
        logging.debug("added ending {}".format(text))

    def set_style(self, paragraph):
        paragraph.runs[0].font.size = self.default_font.size
        paragraph.runs[0].font.rtl = self.default_font.rtl
        paragraph.runs[0].font.name = self.default_font.name
        paragraph.runs[-1].font.size = self.default_font.size
        paragraph.runs[-1].font.rtl = self.default_font.rtl
        paragraph.runs[-1].font.name = self.default_font.name
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    def update_first_page(self):
        index = None
        for i in range(30, 1, -1):
            if self.doc.paragraphs[i].text == FIRST_PAGE_VOLUMES_PARAGRAPH:
                index = i
                break
        if index is not None:
            self.update_text(index, VOLUMES_NAMES_1[self.volume_index - 1])
            self.update_text(index - 1, VOLUMES_NAMES_2[self.volumes_count - 1])
        else:
            exception_message = "נכשל באיתור פסקת זכויות היוצרים ב-30 הפסקאות הראשונות ('{}').".format(FIRST_PAGE_VOLUMES_PARAGRAPH)
            logging.error(exception_message + " לא ניתן לעדכן את העמוד הראשון עם מספר הכרכים.")
            raise ValueError(exception_message)

    def update_text(self, index, new_text):
        runs = self.doc.paragraphs[index].runs
        for r in runs:
            r.clear()
        runs[0].text = new_text
        logging.debug("update paragraph in index {} to {}".format(index, new_text))

    def delete_chapter_beginning_markers(self, elements, chapter_indicator):
        for p in elements:
            if p.is_paragraph and p.block.text == chapter_indicator:
                p.block.text = ''

    def save(self, directory):
        two_digits_volume = str(self.volume_index) if self.volume_index > 9 else "0" + str(self.volume_index)
        volume_path = "{dir}/{docx}/{vol}.{docx}".format(dir=directory, vol=two_digits_volume, docx=DOCX)
        self.doc.save(volume_path)
        self.save_as_doc(volume_path)

    def save_as_doc(self, path):
        if platform.system() == "Windows":
            docx = path.replace("/", "\\")
            doc = docx.replace(DOCX, DOC)
            import win32com.client
            wrd = win32com.client.Dispatch("Word.Application")
            wrd.visible = 0
            wb = wrd.Documents.Open(docx)
            wb.SaveAs(doc, FileFormat=0)
            wb.Close()
            wrd.Quit()
