from docx.opc.exceptions import PackageNotFoundError
from general.book import *
import os
from docx import Document
from general.constants import *
from docx.enum.style import WD_STYLE_TYPE
import logging
from copy import deepcopy
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from general.element import Element
from general.generator import DocxGenerator
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def start_split(input_path, gui_queue, max_chars, min_chars):
    try:
        if not input_path.endswith(DOCX):
            logging.error("input file did not end with 'docx': {} ".format(input_path))
            raise PackageNotFoundError
        output_path = create_split(input_path, max_chars, min_chars)
        gui_queue.put([SUCCESS, output_path])
    except PackageNotFoundError:
        gui_queue.put([FAILURE, "קובץ הספר חייב להיות בפורמט docx"])
    except Exception as error:
        logging.exception(error)
        gui_queue.put([FAILURE, ERROR_MESSAGE.format(error)])
    return


def start_stickers(directory, title, author, volumes, pages, gui_queue):
    try:
        output_path = create_stickers(directory, title, author, volumes, pages)
        gui_queue.put([SUCCESS, output_path])
    except Exception as error:
        logging.exception(error)
        gui_queue.put([FAILURE, ERROR_MESSAGE.format(error)])
    return


def create_split(path, max_chars, min_chars):
    logging.info("start split {} to volumes".format(path))

    origin = Document(path)

    directory = os.path.splitext(path)[0]
    try:
        os.mkdir(directory)
        os.mkdir(os.path.join(directory, DOC))
        os.mkdir(os.path.join(directory, DOCX))

        logging.info("created a directories at {}".format(directory))
    except FileExistsError:
        raise Exception("הספר כבר פוצל בעבר. אנא בדוק את התיקייה {}. במידה ואתה מעוניין לפצל את הספר שוב, אנא מחק את התיקייה ונסה שנית".format(directory))

    book = Book(build_elements_list(origin), max_chars, min_chars)
    volumes = book.volumes
    volumes_count = len(volumes)

    for index, vol in enumerate(volumes, 1):
        new_doc = deepcopy(origin)
        elements = build_elements_list(new_doc)
        formatted = DocxGenerator(new_doc, elements, book.last_gate_paragraph, vol, index, volumes_count)
        formatted.format_volume()
        formatted.save(directory)
    logging.debug("done split to volume, all saved in the file system")

    return directory


def create_stickers(directory, book_title, book_author, num_volumes, num_pages):
    logging.debug("start creating stickers")

    doc = Document()
    style = doc.styles.add_style(STICKER, WD_STYLE_TYPE.PARAGRAPH)
    style.paragraph_format.keep_together = True

    sticker_template = NORMAL_STICKER_TEMPLATE
    if num_pages.strip() != '':
        sticker_template = STUDY_STICKER_TEMPLATE

    for i in range(int(num_volumes)):
        sticker = sticker_template.format(title=book_title,
                                          author=book_author,
                                          total=VOLUMES_NAMES_2[int(num_volumes)-1],
                                          current=VOLUMES_NAMES_1[i],
                                          pages=num_pages)
        p = doc.add_paragraph(text=sticker, style=style)
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    path = "{dir}/{sticker} - {book}.{docx}".format(dir=directory, sticker=STICKER, book=book_title, docx=DOCX)

    doc.save(path)

    logging.debug("done creating stickers, output saved at {}".format(path))
    return path


def download_support():
    doc = Document(SUPPORT_FILE)
    doc.save(SUPPORT_FILE)


def build_elements_list(parent):
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("נכשל במהלך פיצול הספר לאלמנטים של פסקאות וטבלאות")

    elements = []
    for index, child in enumerate(parent_elm.iterchildren()):
        if isinstance(child, CT_P):
            elements.append(Element(True, Paragraph(child, parent)))
        elif isinstance(child, CT_Tbl):
            elements.append(Element(False, Table(child, parent)))
    return elements
