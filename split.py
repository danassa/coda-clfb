from docx import Document
from helpers import Volume
import os
from constants import *


def generate_outputs(path):
    origin = Document(path)

    directory = os.path.splitext(path)[0]
    try:
        os.mkdir(directory)
    except FileExistsError:
        pass

    volumes = split_text_to_volumes(origin)

    volumes_count = len(volumes)
    append_volumes_covers(volumes, volumes_count)
    append_volume_suffix(volumes, volumes_count)
    append_book_suffix(volumes, volumes_count)

    write_volumes(directory, volumes)


def split_text_to_volumes(origin_doc):
    volumes = []
    current_volume = Volume()

    for paragraph in origin_doc.paragraphs:
        paragraph_char_count = len(paragraph.text)
        predicted_volume_count = current_volume.char_count + paragraph_char_count

        if predicted_volume_count >= MAX_CHARS_PER_VOLUME:
            volumes.append(current_volume)
            current_volume = Volume()

        current_volume.add_paragraph(paragraph.text, paragraph_char_count)

    volumes.append(current_volume)

    return volumes


# todo
def append_volumes_covers(volumes, volumes_count):
    prior = 'initial text'
    # volumes[i].paragraphs[0].insert_paragraph_before('prior')


def append_volume_suffix(volumes, volumes_count):
    for x in range(volumes_count - 1):
        doc = volumes[x]
        doc.add_paragraph('', 0)
        last_paragraph = '{}{}'.format(END, VOLUMES_NAMES_1[x])
        doc.add_paragraph(last_paragraph, len(last_paragraph))


#todo verify
def append_book_suffix(volumes, volumes_count):
    doc = volumes[volumes_count - 1]
    doc.add_paragraph('', 0)
    doc.add_paragraph(BOOK_SUFFIX, len(BOOK_SUFFIX))


def write_volumes(directory, outputs):
    for index, volume in enumerate(outputs):
        volume_path = "{dir}/{vol}.{docx}".format(dir=directory, vol=str(index+1), docx=DOCX)
        volume.doc.save(volume_path)

