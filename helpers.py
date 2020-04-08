import platform
from docx import Document
from constants import TEMPLATE
import time


class Volume:
    def __init__(self):
        self.doc = Document(TEMPLATE)
        self.char_count = 0

    def add_paragraph(self, text, char_count):
        self.doc.add_paragraph(text)
        self.char_count += char_count


def get_now():
    return int(round(time.time() * 100))
