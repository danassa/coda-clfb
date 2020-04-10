from constants import *
import time
from docx import Document


class Book:

    def __init__(self, paragraphs, min_chars=MIN_CHARS_PER_VOLUME, max_chars=MAX_CHARS_PER_VOLUME):
        assert min_chars <= max_chars, "minimum characters per volume must be equal or greater than the maximum"

        self.paragraphs = paragraphs
        self.min_chars_per_volume = min_chars
        self.max_chars_per_volume = max_chars

        self.last_gate_paragraph = None
        self.chapters = []
        self.chars = None
        self.volumes = []
        self.find_gate_end()
        self.create_chapters()
        self.split_volumes()

        assert self.last_gate_paragraph is not None, "no copyright paragraph found!"
        assert len(self.chapters) >= 1
        assert len(self.volumes) >= 1

    def find_gate_end(self):
        for index, paragraph in enumerate(self.paragraphs):
            if paragraph.text.startswith(GATE_END_INDICATOR):
                self.last_gate_paragraph = index + 1
                return

    def create_chapters(self):
        total_chars = 0
        chapter_chars = 0
        current_chapter_start_index = self.last_gate_paragraph + 1

        for index, paragraph in enumerate(self.paragraphs):
            if paragraph.text == CHAPTER_INDICATOR:
                if current_chapter_start_index == self.last_gate_paragraph + 1:  # first chapter!
                    current_chapter_start_index = index + 1
                else:
                    new_chapter = Chapter(chapter_chars, current_chapter_start_index, index - 1)
                    self.chapters.append(new_chapter)
                    total_chars += chapter_chars
                    chapter_chars = 0
                    current_chapter_start_index = index + 1
            elif index > self.last_gate_paragraph:
                chapter_chars += len(paragraph.text)

        total_chars += chapter_chars
        last_chapter = Chapter(chapter_chars, current_chapter_start_index, index)
        self.chapters.append(last_chapter)
        self.chars = total_chars

    def split_volumes(self):

        c_index = 0
        v_chars = 0
        v_index = self.last_gate_paragraph + 1
        v_paragraph = None

        while True:
            try:
                chapter = self.chapters[c_index]
            except IndexError:
                if v_chars > 0:
                    self.volumes.append(Volume(v_index, self.chapters[c_index-1].end_paragraph_index, v_paragraph, None))
                break

            if v_chars + chapter.chars > self.max_chars_per_volume and v_chars > self.min_chars_per_volume:
                self.volumes.append(Volume(v_index, chapter.start_paragraph_index - 1, v_paragraph, None))
                v_chars = 0
                v_index = chapter.start_paragraph_index
                v_paragraph = None
            elif v_chars + chapter.chars > self.max_chars_per_volume:
                middle_paragraph_index, chars_left_in_chapter = self.split_chapter(chapter.start_paragraph_index,
                                                                                   chapter.end_paragraph_index,
                                                                                   v_chars + chapter.chars,
                                                                                   self.max_chars_per_volume)
                if middle_paragraph_index != chapter.end_paragraph_index:
                    self.volumes.append(Volume(v_index, middle_paragraph_index, v_paragraph, SPLIT_CHAPTER_END))
                    self.chapters.insert(c_index + 1, Chapter(chars_left_in_chapter, middle_paragraph_index + 1, chapter.end_paragraph_index))
                    self.chapters[c_index].end_paragraph_index = middle_paragraph_index
                    self.chapters[c_index].chars = v_chars + chapter.chars - chars_left_in_chapter
                    v_index = middle_paragraph_index + 1
                    v_paragraph = SPLIT_CHAPTER_START.format(c_index)
                else:
                    self.volumes.append(Volume(v_index, middle_paragraph_index, v_paragraph, None))
                v_chars = 0
                c_index += 1
            elif v_chars + chapter.chars > self.min_chars_per_volume:
                self.volumes.append(Volume(v_index, chapter.end_paragraph_index, v_paragraph, None))
                v_chars = 0
                v_index = chapter.end_paragraph_index + 1
                v_paragraph = None
                c_index += 1
            else:
                v_chars += chapter.chars
                c_index += 1

    # todo אריה במדבר במקום
    def split_chapter(self, first_index, last_index, chars, max_chars):
        current_chars_count = chars
        chars_left_in_chapter = 0
        while current_chars_count > max_chars and last_index > first_index:
            chars_in_paragraph = len(self.paragraphs[last_index].text)
            current_chars_count = current_chars_count - chars_in_paragraph
            last_index -= 1
            chars_left_in_chapter += chars_in_paragraph
        return last_index, chars_left_in_chapter


class Chapter:
    def __init__(self, chars, start, end):
        assert end >= start
        assert chars > 0
        self.chars = chars
        self.start_paragraph_index = start
        self.end_paragraph_index = end


class Volume:
    def __init__(self, start, end, first_paragraph=None, last_paragraph=None):
        self.start_paragraph_index = start
        self.end_paragraph_index = end
        self.first_paragraph = first_paragraph
        self.last_paragraph = last_paragraph
        self.doc = None

    def create_doc(self, path, gate_index, volume_num, last_volume_num):
        self.doc = Document(path)
        self.remove_redundant_paragraph(gate_index)
        self.delete_chapter_beginning_markers()
        self.add_split_chapter_notes(gate_index)
        if volume_num == last_volume_num:
            self.add_book_suffix()
        else:
            self.add_volume_suffix(volume_num)
        self.update_first_page(volume_num, last_volume_num)

    def remove_redundant_paragraph(self, gate_index):
        for p in range(len(self.doc.paragraphs)-1, self.end_paragraph_index, -1):
            self.delete_paragraph(self.doc.paragraphs[p])
        for p in range(self.start_paragraph_index, gate_index, -1):
            self.delete_paragraph(self.doc.paragraphs[p])

    def delete_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def delete_chapter_beginning_markers(self):
        for p in self.doc.paragraphs:
            if p.text == CHAPTER_INDICATOR:
                p.text = ''

    def add_split_chapter_notes(self, gate_index):
        if self.first_paragraph is not None:
            self.doc.paragraphs[gate_index + 1].insert_paragraph_before(self.first_paragraph)
            self.doc.paragraphs[gate_index + 2].insert_paragraph_before('')
        if self.last_paragraph is not None:
            self.doc.add_paragraph('')
            self.doc.add_paragraph(self.last_paragraph)

    def add_volume_suffix(self, vol):
        self.doc.add_paragraph('')
        last_paragraph = '{}{}'.format(END, VOLUMES_NAMES_1[vol])
        self.doc.add_paragraph(last_paragraph)

    def add_book_suffix(self):
        last_non_empty_paragraph = None
        index = len(self.doc.paragraphs) - 1
        while last_non_empty_paragraph is None and index >= 0 :
            if self.doc.paragraphs[index] != '':
                last_non_empty_paragraph = index
            index -= 1

        if self.doc.paragraphs[last_non_empty_paragraph] != BOOK_SUFFIX:
            self.doc.add_paragraph('')
            self.doc.add_paragraph(BOOK_SUFFIX)

    def update_first_page(self, volume_num, last_volume_num):
        index = None
        for i in range(20, 1, -1):
            if self.doc.paragraphs[i].text == FIRST_PAGE_VOLUMES_PARAGRAPH:
                index = i
        if index is not None:
            self.doc.paragraphs[index].text = VOLUMES_NAMES_1[volume_num]
            self.doc.paragraphs[index - 1].text = VOLUMES_NAMES_2[last_volume_num - 1]


def get_now():
    return int(round(time.time() * 100))


def get_elapsed_time_str(start_time):
    elapsed_time = get_now() - start_time
    display_time = '{:02d}:{:02d}'.format((elapsed_time // 100) // 60, (elapsed_time // 100) % 60)
    return display_time
